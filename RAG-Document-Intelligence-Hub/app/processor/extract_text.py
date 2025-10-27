"""
Advanced Text Extraction Module
Multi-format text extraction with OCR support and metadata preservation.
"""

import io
from typing import Dict, Optional, Union
from pathlib import Path

# Core libraries
import pandas as pd
from PyPDF2 import PdfReader

# Optional advanced libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from .detect_format import detect_document_format, DocumentFormat


class ExtractionResult:
    """Container for extraction results with metadata."""
    
    def __init__(self, text: str, metadata: Optional[Dict] = None):
        self.text = text
        self.metadata = metadata or {}
        self.char_count = len(text)
        self.word_count = len(text.split())
        
    def __repr__(self):
        return f"ExtractionResult(chars={self.char_count}, words={self.word_count})"


def extract_text_from_document(
    file_source: Union[str, Path, io.BytesIO],
    filename: Optional[str] = None,
    use_ocr: bool = False
) -> ExtractionResult:
    """
    Extract text from document with automatic format detection.
    
    Args:
        file_source: File path, Path object, or file-like object
        filename: Original filename (required for file objects)
        use_ocr: Enable OCR for scanned PDFs and images
        
    Returns:
        ExtractionResult with text and metadata
        
    Raises:
        ValueError: If format is unsupported
        RuntimeError: If extraction fails
    """
    
    # Detect format
    doc_format = detect_document_format(file_source, filename)
    
    # Route to appropriate extractor
    extractors = {
        'text': _extract_plain_text,
        'pdf': _extract_pdf_text,
        'spreadsheet': _extract_spreadsheet_text,
        'document': _extract_docx_text,
        'ebook': _extract_epub_text,
        'image': _extract_image_text
    }
    
    extractor = extractors.get(doc_format.category)
    if not extractor:
        raise ValueError(f"No extractor available for category: {doc_format.category}")
    
    try:
        text, metadata = extractor(file_source, doc_format, use_ocr)
        
        # Add format info to metadata
        metadata['format'] = doc_format.extension
        metadata['category'] = doc_format.category
        metadata['mime_type'] = doc_format.mime_type
        
        return ExtractionResult(text=text, metadata=metadata)
        
    except Exception as e:
        raise RuntimeError(f"Extraction failed for {doc_format.extension}: {str(e)}")


# ========== Format-specific extractors ==========

def _extract_plain_text(file_source, doc_format: DocumentFormat, use_ocr: bool):
    """Extract from plain text files (.txt, .md, .rst, etc.)"""
    
    if isinstance(file_source, (str, Path)):
        with open(file_source, 'r', encoding='utf-8') as f:
            text = f.read()
    else:
        text = file_source.read().decode('utf-8')
    
    metadata = {
        'lines': text.count('\n') + 1
    }
    
    return text, metadata


def _extract_pdf_text(file_source, doc_format: DocumentFormat, use_ocr: bool):
    """Extract from PDF files with fallback options."""
    
    # Try pdfplumber first (better quality)
    if PDFPLUMBER_AVAILABLE:
        try:
            return _extract_pdf_with_pdfplumber(file_source)
        except Exception:
            pass  # Fall back to PyPDF2
    
    # Fallback to PyPDF2
    return _extract_pdf_with_pypdf2(file_source, use_ocr)


def _extract_pdf_with_pdfplumber(file_source):
    """Extract PDF using pdfplumber (better layout preservation)."""
    
    import pdfplumber
    
    if isinstance(file_source, (str, Path)):
        pdf = pdfplumber.open(file_source)
    else:
        pdf = pdfplumber.open(file_source)
    
    pages_text = []
    for page in pdf.pages:
        text = page.extract_text()
        if text:
            pages_text.append(text)
    
    pdf.close()
    
    full_text = "\n\n".join(pages_text)
    
    metadata = {
        'pages': len(pages_text),
        'extraction_method': 'pdfplumber'
    }
    
    return full_text, metadata


def _extract_pdf_with_pypdf2(file_source, use_ocr: bool):
    """Extract PDF using PyPDF2."""
    
    if isinstance(file_source, (str, Path)):
        with open(file_source, 'rb') as f:
            reader = PdfReader(f)
            pages_text = [
                page.extract_text() 
                for page in reader.pages 
                if page.extract_text()
            ]
    else:
        reader = PdfReader(file_source)
        pages_text = [
            page.extract_text() 
            for page in reader.pages 
            if page.extract_text()
        ]
    
    full_text = "\n\n".join(pages_text)
    
    # Check if OCR might be needed
    if len(full_text.strip()) < 100 and use_ocr and OCR_AVAILABLE:
        # TODO: Implement OCR fallback for scanned PDFs
        pass
    
    metadata = {
        'pages': len(pages_text),
        'extraction_method': 'pypdf2'
    }
    
    return full_text, metadata


def _extract_spreadsheet_text(file_source, doc_format: DocumentFormat, use_ocr: bool):
    """Extract from spreadsheet files (.csv, .xlsx, etc.)"""
    
    if doc_format.extension == 'csv':
        df = pd.read_csv(file_source)
    elif doc_format.extension in ['xlsx', 'xls']:
        df = pd.read_excel(file_source)
    elif doc_format.extension == 'tsv':
        df = pd.read_csv(file_source, sep='\t')
    else:
        raise ValueError(f"Unsupported spreadsheet format: {doc_format.extension}")
    
    # Convert to readable text format
    text = df.to_string(index=False)
    
    metadata = {
        'rows': len(df),
        'columns': len(df.columns),
        'column_names': list(df.columns)
    }
    
    return text, metadata


def _extract_docx_text(file_source, doc_format: DocumentFormat, use_ocr: bool):
    """Extract from Word documents (.docx)"""
    
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
    
    if isinstance(file_source, (str, Path)):
        doc = DocxDocument(file_source)
    else:
        doc = DocxDocument(file_source)
    
    # Extract paragraphs
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    text = "\n\n".join(paragraphs)
    
    metadata = {
        'paragraphs': len(paragraphs),
        'extraction_method': 'python-docx'
    }
    
    return text, metadata


def _extract_epub_text(file_source, doc_format: DocumentFormat, use_ocr: bool):
    """Extract from EPUB ebooks."""
    
    if not EPUB_AVAILABLE:
        raise RuntimeError("ebooklib not installed. Install with: pip install ebooklib beautifulsoup4")
    
    if isinstance(file_source, (str, Path)):
        book = epub.read_epub(file_source)
    else:
        # Save to temp file for ebooklib
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.epub') as tmp:
            tmp.write(file_source.read())
            tmp_path = tmp.name
        book = epub.read_epub(tmp_path)
    
    chapters = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text = soup.get_text()
            if text.strip():
                chapters.append(text)
    
    full_text = "\n\n".join(chapters)
    
    metadata = {
        'chapters': len(chapters),
        'extraction_method': 'ebooklib'
    }
    
    return full_text, metadata


def _extract_image_text(file_source, doc_format: DocumentFormat, use_ocr: bool):
    """Extract text from images using OCR."""
    
    if not OCR_AVAILABLE:
        raise RuntimeError("OCR not available. Install with: pip install pytesseract pillow")
    
    if not use_ocr:
        raise ValueError("OCR is disabled. Set use_ocr=True to extract text from images.")
    
    if isinstance(file_source, (str, Path)):
        image = Image.open(file_source)
    else:
        image = Image.open(file_source)
    
    # Perform OCR
    text = pytesseract.image_to_string(image)
    
    metadata = {
        'image_size': image.size,
        'extraction_method': 'tesseract-ocr'
    }
    
    return text, metadata
